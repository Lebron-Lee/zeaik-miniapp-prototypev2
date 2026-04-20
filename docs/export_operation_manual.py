from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

BASE = Path('/home/ubuntu/zeaik-miniapp-prototype/docs')
MD_PATH = BASE / 'zeaik_miniapp_operation_manual.md'
DOCX_PATH = BASE / '智爱客小程序完整操作手册.docx'


def resolve_image_path(raw: str) -> Path:
    return (BASE / raw).resolve()


def add_table(doc: Document, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    header = table.rows[0].cells
    for i, cell in enumerate(rows[0]):
        header[i].text = cell.strip()
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            cells[i].text = cell.strip()
    doc.add_paragraph('')


def flush_table(doc: Document, table_lines):
    if not table_lines:
        return
    rows = []
    for idx, line in enumerate(table_lines):
        parts = [p.strip() for p in line.strip().strip('|').split('|')]
        if idx == 1 and all(set(p) <= set('-: ') for p in parts):
            continue
        rows.append(parts)
    if rows:
        add_table(doc, rows)


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    lines = MD_PATH.read_text(encoding='utf-8').splitlines()
    table_lines = []

    for raw in lines:
        line = raw.rstrip('\n')

        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines.append(line)
            continue
        else:
            flush_table(doc, table_lines)
            table_lines = []

        if not line.strip():
            doc.add_paragraph('')
            continue

        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            caption, img_rel = img_match.groups()
            img_path = resolve_image_path(img_rel)
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(3.1))
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph(f'[缺失图片] {caption} -> {img_rel}')
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.italic = True
        else:
            clean = re.sub(r'\[(\d+)\]', r'\1', line)
            clean = clean.replace('**', '')
            doc.add_paragraph(clean)

    flush_table(doc, table_lines)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == '__main__':
    main()
